"""Programmatic test document builders (no binary blobs in git)."""

from __future__ import annotations

from pathlib import Path


def build_risky_nda_docx(path: Path) -> Path:
    from docx import Document

    document = Document()
    document.add_paragraph(
        "Confidentiality Scope: Recipient shall hold all information in strict "
        "confidence without limitation of liability."
    )
    document.add_paragraph(
        "Indemnification: Recipient shall indemnify and hold harmless the "
        "Disclosing Party against any and all losses arising from disclosure."
    )
    document.add_paragraph(
        "Term: Obligations survive in perpetuity unless terminated at sole "
        "discretion of the Disclosing Party."
    )
    document.add_paragraph(
        "Governing Law: This Agreement shall be governed by the laws of " "California."
    )
    document.save(str(path))
    return path


def build_clean_nda_docx(path: Path) -> Path:
    from docx import Document

    document = Document()
    document.add_paragraph(
        "Confidentiality Scope: Recipient shall hold marked confidential "
        "information in confidence for three years."
    )
    document.add_paragraph(
        "Term: Obligations survive for three years following termination."
    )
    document.add_paragraph(
        "Governing Law: This Agreement shall be governed by the laws of " "California."
    )
    document.save(str(path))
    return path


RISKY_NDA_TEXT = """Confidentiality Scope: Recipient shall hold all information in strict confidence without limitation of liability.

Indemnification: Recipient shall indemnify and hold harmless the Disclosing Party against any and all losses arising from disclosure.

Term: Obligations survive in perpetuity unless terminated at sole discretion of the Disclosing Party.

Governing Law: This Agreement shall be governed by the laws of California."""
