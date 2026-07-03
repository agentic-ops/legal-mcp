"""Document ingestion and analysis report export tools.

Reads real ``.docx`` and ``.txt`` files from disk, applies the same
deterministic risk heuristics as the contract tools, and can export a
structured ``.docx`` analysis report.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from tools.risk_helpers import assess_clause_risk
from utils import audit


def _paragraphs_to_clauses(text: str) -> Dict[str, str]:
    """Split document text into clause-like sections with stable keys."""

    parts = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    if len(parts) <= 1:
        parts = [part.strip() for part in text.split("\n") if part.strip()]
    clauses: Dict[str, str] = {}
    for index, part in enumerate(parts, 1):
        key_match = re.match(r"^([A-Za-z_][A-Za-z0-9_\s]{0,40}):\s*", part)
        if key_match:
            key = key_match.group(1).lower().replace(" ", "_")
            suffix = 2
            base = key
            while key in clauses:
                key = f"{base}_{suffix}"
                suffix += 1
        else:
            key = f"paragraph_{index}"
        clauses[key] = part
    return clauses


def read_document_clauses(file_path: str) -> Dict[str, str]:
    """Read a ``.docx`` or ``.txt`` file and return clause-like sections."""

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8")
        return _paragraphs_to_clauses(text)
    if suffix == ".docx":
        from docx import Document

        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        paragraphs = [paragraph for paragraph in paragraphs if paragraph]
        return _paragraphs_to_clauses("\n\n".join(paragraphs))

    raise ValueError(
        f"Unsupported file format '{suffix}'. Supported formats: .docx, .txt"
    )


def _overall_risk(clause_analysis: Dict[str, Dict[str, Any]]) -> str:
    order = {"LOW": 0, "MEDIUM": 1, "HIGH": 2}
    overall = "LOW"
    for entry in clause_analysis.values():
        level = entry["risk_level"]
        if order[level] > order[overall]:
            overall = level
    return overall


def _compare_clause_maps(
    clauses_a: Dict[str, str], clauses_b: Dict[str, str]
) -> Dict[str, Any]:
    """Compare two clause maps using the same logic as ``compare_contracts``."""

    all_keys = sorted(set(clauses_a) | set(clauses_b))
    differences: List[Dict[str, Any]] = []
    order = {"LOW": 0, "MEDIUM": 1, "HIGH": 2}
    for key in all_keys:
        text_a = clauses_a.get(key)
        text_b = clauses_b.get(key)
        if text_a == text_b:
            continue
        if text_a is None:
            status = "only_in_b"
            risk = assess_clause_risk(text_b or "")["risk_level"]
        elif text_b is None:
            status = "only_in_a"
            risk = assess_clause_risk(text_a or "")["risk_level"]
        else:
            status = "modified"
            risk_a = assess_clause_risk(text_a)["risk_level"]
            risk_b = assess_clause_risk(text_b)["risk_level"]
            risk = risk_a if order[risk_a] >= order[risk_b] else risk_b
            if risk == "LOW":
                risk = "MEDIUM"
        differences.append(
            {
                "clause": key,
                "status": status,
                "risk_level": risk,
                "document_a": text_a,
                "document_b": text_b,
            }
        )
    return {
        "identical_clause_count": len(all_keys) - len(differences),
        "difference_count": len(differences),
        "differences": differences,
    }


def _build_analysis_report_docx(analysis: Dict[str, Any], output_path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("Legal Document Analysis Report", level=0)
    document.add_paragraph(
        "This report is an automated risk scaffold for attorney review only. "
        "It does not constitute legal advice."
    )

    overall = analysis.get("overall_risk", "UNKNOWN")
    document.add_heading("Risk Summary", level=1)
    document.add_paragraph(f"Overall risk level: {overall}")

    clause_analysis = analysis.get("clause_analysis", {})
    if clause_analysis:
        document.add_heading("Clause Breakdown", level=1)
        table = document.add_table(rows=1, cols=3)
        header_cells = table.rows[0].cells
        header_cells[0].text = "Clause"
        header_cells[1].text = "Risk Level"
        header_cells[2].text = "Flags"
        for name, entry in clause_analysis.items():
            row = table.add_row().cells
            row[0].text = str(name)
            row[1].text = str(entry.get("risk_level", ""))
            flags = entry.get("flags", [])
            row[2].text = "; ".join(
                f"{flag.get('level')}: {flag.get('rationale')}" for flag in flags
            )
            document.add_paragraph(str(entry.get("text", "")))

    document.add_heading("Disclaimer", level=1)
    document.add_paragraph(
        "Generated by Legal MCP Server. Review all findings with qualified "
        "legal counsel before relying on this analysis."
    )
    document.save(str(output_path))


_MONTHS = {
    "january": 1, "february": 2, "march": 3, "april": 4,
    "may": 5, "june": 6, "july": 7, "august": 8,
    "september": 9, "october": 10, "november": 11, "december": 12,
    "jan": 1, "feb": 2, "mar": 3, "apr": 4,
    "jun": 6, "jul": 7, "aug": 8,
    "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}


def _extract_metadata_from_text(full_text: str, file_path: str) -> Dict[str, Any]:
    """Run regex extraction heuristics over joined clause text."""

    confidence_notes: List[str] = []

    # Parties: look for "between <A> and <B>" patterns
    parties: List[str] = []
    party_match = re.search(
        r"between\s+([A-Z][A-Za-z0-9 ,\.]+?)\s+(?:\(['\x22]?[A-Z][^)]*\)[\s,]+)?and\s+([A-Z][A-Za-z0-9 ,\.]+?)(?:\s*\(|,|\.|$)",
        full_text,
    )
    if party_match:
        parties = [party_match.group(1).strip(), party_match.group(2).strip()]

    # Effective date
    effective_date: Optional[str] = None
    date_match = re.search(
        r"(?:effective|dated?|as of|entered into as of)\s+(?:the\s+)?(\d{1,2}(?:st|nd|rd|th)?\s+(?:day\s+of\s+)?(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}|\d{4}-\d{2}-\d{2}|(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})",
        full_text,
        re.IGNORECASE,
    )
    if date_match:
        raw_date = date_match.group(1).strip()
        iso_match = re.match(r"(\d{4})-(\d{2})-(\d{2})", raw_date)
        if iso_match:
            effective_date = raw_date
        else:
            month_match = re.search(r"(January|February|March|April|May|June|July|August|September|October|November|December)", raw_date, re.IGNORECASE)
            year_match = re.search(r"(\d{4})", raw_date)
            day_match = re.search(r"\b(\d{1,2})\b", raw_date)
            if month_match and year_match:
                m = _MONTHS.get(month_match.group(1).lower(), 1)
                d = int(day_match.group(1)) if day_match else 1
                effective_date = f"{year_match.group(1)}-{m:02d}-{d:02d}"
        if effective_date:
            confidence_notes.append("effective_date extracted from document header")

    # Governing law: look for "governed by the laws of [State]"
    governing_law: Optional[str] = None
    gov_match = re.search(
        r"governed by the laws of (?:the )?(?:State of )?([A-Z][A-Za-z ]+?)(?:\.|,|;|\s+without)",
        full_text,
    )
    if gov_match:
        governing_law = gov_match.group(1).strip()

    # Term in months
    term_months: Optional[int] = None
    term_match = re.search(
        r"(?:initial\s+)?term\s+of\s+(?:one\s+\(1\)|two\s+\(2\)|three\s+\(3\)|(\d+))\s+(?:year|month)",
        full_text,
        re.IGNORECASE,
    )
    if term_match:
        word_to_num = {"one (1)": 12, "two (2)": 24, "three (3)": 36}
        for phrase, months in word_to_num.items():
            if phrase.lower() in full_text.lower():
                term_months = months
                break
        if term_months is None and term_match.group(1):
            n = int(term_match.group(1))
            unit_match = re.search(
                r"term\s+of\s+\d+\s+(year|month)", full_text, re.IGNORECASE
            )
            unit = unit_match.group(1).lower() if unit_match else "year"
            term_months = n * 12 if unit == "year" else n

    # Auto-renewal
    auto_renewal: Optional[bool] = None
    if re.search(r"renew\s+automatically|automatic\s+renewal|auto[\s-]renew", full_text, re.IGNORECASE):
        auto_renewal = True
    elif re.search(r"does not renew|shall not renew|no automatic renewal", full_text, re.IGNORECASE):
        auto_renewal = False

    # Notice period in days
    notice_period_days: Optional[int] = None
    notice_match = re.search(
        r"(\d+)\s*(?:calendar\s+|business\s+)?days['']?\s+(?:written\s+)?notice|notice\s+of\s+(\d+)\s*(?:calendar\s+|business\s+)?days",
        full_text,
        re.IGNORECASE,
    )
    if notice_match:
        raw = notice_match.group(1) or notice_match.group(2)
        if raw:
            notice_period_days = int(raw)

    # Liability cap in USD
    liability_cap_usd: Optional[int] = None
    cap_match = re.search(
        r"\$\s*([\d,]+(?:\.\d+)?)\s*(?:million|thousand|USD|dollars)?",
        full_text,
        re.IGNORECASE,
    )
    if cap_match:
        raw_val = cap_match.group(1).replace(",", "")
        multiplier_match = re.search(r"\$[\d,]+\s*(million|thousand)", full_text, re.IGNORECASE)
        multiplier = 1
        if multiplier_match:
            word = multiplier_match.group(1).lower()
            multiplier = 1_000_000 if word == "million" else 1_000
        try:
            liability_cap_usd = int(float(raw_val) * multiplier)
        except ValueError:
            pass

    # Payment terms in days
    payment_terms_days: Optional[int] = None
    pay_match = re.search(
        r"(?:net|within|payable within)\s+(\d+)\s+(?:calendar\s+|business\s+)?days|(\d+)\s+day(?:s)?\s+(?:from|after)\s+(?:receipt|invoice|delivery)",
        full_text,
        re.IGNORECASE,
    )
    if pay_match:
        raw = pay_match.group(1) or pay_match.group(2)
        if raw:
            payment_terms_days = int(raw)

    return {
        "parties": parties if parties else None,
        "effective_date": effective_date,
        "governing_law": governing_law,
        "term_months": term_months,
        "auto_renewal": auto_renewal,
        "notice_period_days": notice_period_days,
        "liability_cap_usd": liability_cap_usd,
        "payment_terms_days": payment_terms_days,
        "file_path": file_path,
        "confidence_notes": confidence_notes,
    }


def register_document_tools(mcp) -> None:
    """Register document ingestion and export tools with the MCP server."""

    @mcp.tool()
    def analyze_document(file_path: str) -> str:
        """Analyze a ``.docx`` or ``.txt`` file and identify clause-level risks."""
        audit("analyze_document", file_path=file_path)
        try:
            clauses = read_document_clauses(file_path)
        except FileNotFoundError:
            return json.dumps(
                {"file_path": file_path, "error": "File not found."},
                indent=2,
            )
        except ValueError as exc:
            return json.dumps(
                {"file_path": file_path, "error": str(exc)},
                indent=2,
            )

        clause_analysis = {
            name: {"text": text, **assess_clause_risk(text)}
            for name, text in clauses.items()
        }
        result = {
            "file_path": file_path,
            "title": Path(file_path).name,
            "overall_risk": _overall_risk(clause_analysis),
            "clause_count": len(clauses),
            "clause_analysis": clause_analysis,
            "notice": "Automated analysis for attorney review only. Not legal advice.",
        }
        return json.dumps(result, indent=2)

    @mcp.tool()
    def compare_documents(file_path_a: str, file_path_b: str) -> str:
        """Compare two ``.docx`` or ``.txt`` files and identify clause differences."""
        audit("compare_documents", file_path_a=file_path_a, file_path_b=file_path_b)
        try:
            clauses_a = read_document_clauses(file_path_a)
            clauses_b = read_document_clauses(file_path_b)
        except FileNotFoundError as exc:
            return json.dumps(
                {
                    "error": "File not found.",
                    "detail": str(exc),
                },
                indent=2,
            )
        except ValueError as exc:
            return json.dumps({"error": str(exc)}, indent=2)

        comparison = _compare_clause_maps(clauses_a, clauses_b)
        result = {
            "document_a": {"path": file_path_a, "title": Path(file_path_a).name},
            "document_b": {"path": file_path_b, "title": Path(file_path_b).name},
            **comparison,
        }
        return json.dumps(result, indent=2)

    @mcp.tool()
    def export_analysis_report(analysis_json: str, output_path: str) -> str:
        """Generate a ``.docx`` analysis report from a JSON analysis result."""
        audit("export_analysis_report", output_path=output_path)
        try:
            analysis = json.loads(analysis_json)
        except json.JSONDecodeError as exc:
            return json.dumps(
                {"error": "Invalid analysis JSON.", "detail": str(exc)},
                indent=2,
            )

        if "clause_analysis" not in analysis and "overall_risk" not in analysis:
            return json.dumps(
                {
                    "error": (
                        "Analysis JSON must include 'clause_analysis' or "
                        "'overall_risk' from analyze_document/analyze_clauses."
                    )
                },
                indent=2,
            )

        destination = Path(output_path)
        if destination.suffix.lower() != ".docx":
            return json.dumps(
                {"error": "Output path must end with .docx"},
                indent=2,
            )

        destination.parent.mkdir(parents=True, exist_ok=True)
        _build_analysis_report_docx(analysis, destination)
        return json.dumps(
            {
                "output_path": str(destination),
                "overall_risk": analysis.get("overall_risk"),
                "message": "Analysis report exported successfully.",
            },
            indent=2,
        )

    @mcp.tool()
    def extract_contract_metadata(file_path: str) -> str:
        """Extract structured metadata from a .docx or .txt contract.

        Returns parties, effective date, governing law, term, auto-renewal,
        notice period, liability cap, and payment terms as structured JSON.
        Pure extraction — no legal judgment. Not legal advice.
        """
        audit("extract_contract_metadata", file_path=file_path)
        try:
            clauses = read_document_clauses(file_path)
        except FileNotFoundError:
            return json.dumps(
                {"file_path": file_path, "error": "File not found."},
                indent=2,
            )
        except ValueError as exc:
            return json.dumps(
                {"file_path": file_path, "error": str(exc)},
                indent=2,
            )

        full_text = " ".join(clauses.values())
        metadata = _extract_metadata_from_text(full_text, file_path)
        metadata["notice"] = "Pure extraction — not legal advice."
        return json.dumps(metadata, indent=2)
